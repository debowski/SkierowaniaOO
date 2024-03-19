import os
import subprocess
import tkinter as tk
import tkinter.ttk as ttk
import ttkbootstrap as ttkb
from tkinter import StringVar
from tkinter import filedialog
import pandas as pd
from docxtpl import DocxTemplate
import docx
from docx2pdf import convert


class App:

    def __init__(self):
        self.root = ttkb.Window(themename="darkly")
        self.root.title("Skierowania 0.27")
        self.root.grid()
        self.root.columnconfigure(0, weight=0, minsize=500)
        self.root.columnconfigure(1, weight=1, minsize=400)
        self.root.rowconfigure(0, weight=1)
        self.dodaj_widzety()
        self.credits()

    def dodaj_widzety(self):
        self.frame = ttkb.Frame(self.root)
        self.frame.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)

        self.frame2 = ttkb.Frame(self.root, bootstyle="success")
        self.frame2.grid(row=0, column=1, sticky="nsew", padx=10, pady=10)

        self.frame.columnconfigure(0, weight=1)
        self.frame.columnconfigure(1, weight=1)
        self.frame.rowconfigure(0, weight=0)
        self.frame.rowconfigure(9, weight=1)

        self.frame2.columnconfigure(0, weight=1)
        self.frame2.rowconfigure(0, weight=1)

        self.plik = ""

        self.btn_wyb_plik = ttkb.Button(
            self.frame, text="Wybierz plik", bootstyle="info", command=self.otwarcie_pliku)
        self.btn_wyb_plik.grid(row=0, column=0, sticky="nsew",
                               columnspan=3, padx=5, pady=5)

        # Wybieranie klasy

        self.var = StringVar()

        self.radio1 = ttkb.Radiobutton(
            self.frame, text="Klasa 1", variable=self.var, value="1", bootstyle="info-outline-toolbutton", command=self.set_lista_zawodow_1)
        self.radio2 = ttkb.Radiobutton(
            self.frame, text="Klasa 2", variable=self.var, value="2", bootstyle="info-outline-toolbutton", command=self.set_lista_zawodow_2)
        self.radio3 = ttkb.Radiobutton(
            self.frame, text="Klasa 3", variable=self.var, value="3", bootstyle="info-outline-toolbutton", command=self.set_lista_zawodow_3)
        self.radio1.grid(row=1, column=0, sticky="nsew", padx=5, pady=5)
        self.radio2.grid(row=1, column=1, sticky="nsew", padx=5, pady=5)
        self.radio3.grid(row=1, column=2, sticky="nsew", padx=5, pady=5)

        self.radiobuttons = []
        self.lista_zawodow = StringVar()

        self.zawody = ('Sprzedawca',)
        self.zawody1 = ('Wybierz mądrze...',)
        self.zawody2 = ('Wybierz mądrze...',)
        self.zawody3 = ('Wybierz mądrze...',)

        # wstaw combobox z wyborem zawodow

        # Set the default value for the combobox to be the first item in the list of values.
        self.combo_current_var = tk.StringVar()
        self.combobox = ttkb.Combobox(
            self.frame, values=self.zawody, textvariable=self.combo_current_var, bootstyle="success")
        self.combobox.grid(row=2, column=0, sticky="nsew",
                           padx=5, pady=5, columnspan=3)
        self.combobox.configure(state='readonly')
        self.combobox.set(self.zawody[0])
        self.combobox.current(0)

        self.combobox.bind("<<ComboboxSelected>>", self.wypisanie_osob)

        self.lab_data_wystawienia = ttkb.Label(
            self.frame, text="Data wystawienia")
        self.lab_data_wystawienia.grid(
            row=3, column=0, sticky="nsew", padx=5, pady=5)

        self.data_wystawienia = ttkb.DateEntry(
            self.frame, firstweekday=0)
        self.data_wystawienia.grid(
            row=3, column=1, sticky="nsew", padx=5, pady=5, columnspan=2)

        self.lab_data_rozpoczecia = ttkb.Label(
            self.frame, text="Data rozpoczęcia")
        self.lab_data_rozpoczecia.grid(
            row=4, column=0, sticky="nsew", padx=5, pady=5)

        self.data_rozpoczecia = ttkb.DateEntry(
            self.frame)
        self.data_rozpoczecia.grid(
            row=4, column=1, sticky="nsew", padx=5, pady=5, columnspan=2)

        self.lab_data_zakonczenia = ttkb.Label(
            self.frame, text="Data zakończenia")
        self.lab_data_zakonczenia.grid(
            row=5, column=0, sticky="nsew", padx=5, pady=5)

        self.data_zakonczenia = ttkb.DateEntry(
            self.frame)
        self.data_zakonczenia.grid(
            row=5, column=1, sticky="nsew", padx=5, pady=5, columnspan=2)

        self.label_godzina_rozpoczecia = ttkb.Label(
            self.frame, text="Godzina rozpoczęcia")
        self.label_godzina_rozpoczecia.grid(
            row=6, column=0, sticky="nsew", padx=5, pady=5)

        self.godzina_rozpoczecia = ttkb.Spinbox(
            self.frame, from_=0, to=23, justify="center", format="%02.0f")
        self.godzina_rozpoczecia.grid(
            row=6, column=1, sticky="nsew", padx=5, pady=5)
        self.godzina_rozpoczecia.insert(0, "08")

        self.minuty_rozpoczecia = ttkb.Spinbox(
            self.frame, from_=0, to=59, justify="center", format="%02.0f")
        self.minuty_rozpoczecia.grid(
            row=6, column=2, sticky="nsew", padx=5, pady=5)
        self.minuty_rozpoczecia.insert(0, "00")

        self.btn_utworz_wykaz = ttkb.Button(
            self.frame, text="Utwórz wykaz", bootstyle="success", command=self.utworz_wykaz)
        self.btn_utworz_wykaz.grid(
            row=7, column=0, sticky="nsew", padx=5, pady=5)

        self.btn_utworz_wykaz_pdf = ttkb.Button(
            self.frame, text="Konwersja do PDF", bootstyle="info", command=self.utworz_wykaz_pdf)
        self.btn_utworz_wykaz_pdf.grid(
            row=7, column=1, sticky="nsew", padx=5, pady=5)

        self.btn_otworz_folder_wykaz = ttkb.Button(
            self.frame, text="Otwórz folder wykaz", command=self.otworz_folder_wykaz)
        self.btn_otworz_folder_wykaz.grid(
            row=7, column=2, sticky="nsew", padx=5, pady=5, columnspan=1)

        self.btn_utworz_skierowania = ttkb.Button(
            self.frame, text="Utwórz skierowania", bootstyle="success", command=self.utworz_skierowania)
        self.btn_utworz_skierowania.grid(
            row=8, column=0, sticky="nsew", padx=5, pady=5)

        self.btn_utworz_skierowania_pdf = ttkb.Button(
            self.frame, text="Konwersja do PDF", bootstyle="info", command=self.utworz_skierowania_pdf)
        self.btn_utworz_skierowania_pdf.grid(
            row=8, column=1, sticky="nsew", padx=5, pady=5)

        self.btn_otworz_folder_skierowania = ttkb.Button(
            self.frame, text="Otwórz folder skierowania", command=self.otworz_folder_skierowania)
        self.btn_otworz_folder_skierowania.grid(
            row=8, column=2, sticky="nsew", padx=5, pady=5, columnspan=1)

        self.wynik = ttkb.Label(self.frame, text="Wynik",
                                bootstyle="inverse-dark")
        self.wynik.grid(row=9, column=0, sticky="sew",
                        padx=5, pady=5, columnspan=3)

        # self.separator = ttkb.Separator(self.frame, orient="horizontal", bootstyle="success")
        # self.separator.grid(row=9, column=0, sticky="sew", padx=5, pady=5, columnspan=3)

        # frame2 - przyciski
        self.pole_tekstowe = tk.Text(self.frame2)
        self.pole_tekstowe.grid(
            row=0, column=0,  padx=1, pady=1, sticky="nsew")

    def set_lista_zawodow_1(self) -> None:
        self.combobox['values'] = list(self.zawody1)
        self.wypisanie_osob()

    def set_lista_zawodow_2(self) -> None:
        self.combobox['values'] = list(self.zawody2)
        self.wypisanie_osob()

    def set_lista_zawodow_3(self) -> None:
        self.combobox['values'] = list(self.zawody3)
        self.wypisanie_osob()

    def utworz_lz(self):
        df = pd.read_excel(self.plik)
        # print(df.head)

        # Ekstrakcja liczby z kolumny 'Dane oddziału'
        df['Oddział'] = df['Dane oddziału'].str.extract(
            '(\d+)', expand=False).astype(int)

        # Utworzenie trzech zbiorów na podstawie wartości kolumny 'Oddział'
        self.zawody1 = set(df[df['Oddział'] == 1]
                           ['Specjalność/Zawód'].tolist())
        self.zawody2 = set(df[df['Oddział'] == 2]
                           ['Specjalność/Zawód'].tolist())
        self.zawody3 = set(df[df['Oddział'] == 3]
                           ['Specjalność/Zawód'].tolist())

    # Funkcja sprawdzająca strukturę kolumn w pliku xlsx
    def check_columns(self, file_path):
        expected_columns = ['PESEL', 'Data urodzenia', 'Specjalność/Zawód', 'Miejsce urodzenia', 'Imię', 'Dane oddziału', 'Nazwisko']

        try:
            df = pd.read_excel(file_path)
            columns = df.columns.tolist()

            if set(expected_columns).issubset(set(columns)):
                print("Plik ma poprawną strukturę kolumn.") 
                self.btn_wyb_plik.configure(bootstyle="success")                 
                self.btn_utworz_wykaz.configure(state="normal")
                self.btn_utworz_skierowania.configure(state="normal")



                return True
            else:
                self.btn_wyb_plik.configure(text="Niepoprawne dane")
                self.btn_wyb_plik.configure(bootstyle="danger")

                self.btn_utworz_wykaz.configure(state="disabled")
                self.btn_utworz_skierowania.configure(state="disabled")

                print("Plik nie zawiera wszystkich oczekiwanych kolumn.")



                return False

        except Exception as e:
            print("Błąd podczas sprawdzania pliku:", e)
            return False
   
    def otwarcie_pliku(self):
        filetypes = (
            ('Arkusze', '*.xlsx'),
            ('All files', '*.*')
        )


        print(os.getcwd())  

        self.plik = filedialog.askopenfilename(title='Wybierz plik',
                                               initialdir='..\\Data',
                                               filetypes=filetypes)


        if self.plik:  # Sprawdzamy, czy plik został wybrany
            if self.check_columns(self.plik):  # Wywołujemy funkcję sprawdzającą kolumny
                self.btn_wyb_plik.configure(text=self.plik)
                self.utworz_lz()
                self.radio1.invoke()
            else:
                # Jeśli funkcja check_columns zwróci False, to znaczy że struktura kolumn jest nieprawidłowa
                # Tutaj możesz dodać kod obsługi sytuacji, gdy struktura kolumn jest nieprawidłowa
                

                
                print("Nieprawidłowa struktura kolumn w pliku.")
        else:
            # Jeśli użytkownik nie wybrał pliku
            print("Nie wybrano pliku.")

    def wypisanie_osob(self, event=None):

        if self.plik == "":
            self.brak_pliku()
            return
        else:

            df = pd.read_excel(open(self.plik, "rb"))
            filtered_df = df[df["Dane oddziału"].str.contains(self.var.get(
            ), case=False) & df['Specjalność/Zawód'].str.contains(self.combobox.get(), case=False)]
            tekst = ""
            numer = 1

            for linia in range(filtered_df.shape[0]):
                rekord = filtered_df.iloc[linia].to_dict()

                tekst = tekst + str(numer) + ". " + \
                    rekord['Imię'] + " " + rekord['Nazwisko'] + "\n"

                numer = numer+1

            # wstawianie listy uczniów do ramki prawej
            self.pole_tekstowe.delete(1.0, tk.END)
            self.pole_tekstowe.insert(tk.END, tekst)

    def brak_pliku(self):
        self.pole_tekstowe.delete(1.0, tk.END)
        self.pole_tekstowe.insert(tk.END, "Nie wybrano pliku")

    def symbolZawodu(self, specjalnosc) -> str:

        zawody_dict = {
            "Administrator produkcji filmowej i telewizyjnej (Wprowadzono na podstawie rozporządzenia z dnia 18 stycznia 2023 r.)": "343919",
            "Animator rynku książki (Wprowadzono na podstawie rozporządzenia z dnia 18 stycznia 2023 r.)": "343305",
            "Asystent kierownika produkcji filmowej i telewizyjnej": "343902",
            "Asystent osoby niepełnosprawnej": "341201",
            "Asystentka stomatologiczna": "325101",
            "Automatyk": "731107",
            "Betoniarz-zbrojarz": "711402",
            "Blacharz": "721301",
            "Blacharz samochodowy": "721306",
            "Cieśla": "711501",
            "Cukiernik": "751201",
            "Dekarz": "712101",
            "Drukarz fleksograficzny": "732209",
            "Drukarz offsetowy": "732210",
            "Elektromechanik": "741201",
            "Elektromechanik pojazdów samochodowych": "741203",
            "Elektronik": "742117",
            "Elektryk": "741103",
            "Florysta": "343203",
            "Fotograf": "343101",
            "Fryzjer": "514101",
            "Garbarz skór": "753501",
            "Górnik eksploatacji otworowej": "811301",
            "Górnik eksploatacji podziemnej": "811101",
            "Górnik odkrywkowej eksploatacji złóż": "811102",
            "Górnik podziemnej eksploatacji kopalin innych niż węgiel kamienny": "811112",
            "Higienistka stomatologiczna": "325102",
            "Jeździec": "516408",
            "Kaletnik": "753702",
            "Kamieniarz": "711301",
            "Kelner": "513101",
            "Kierowca mechanik": "832201",
            "Kominiarz": "713303",
            "Koszykarz-plecionkarz": "731702",
            "Kowal": "722101",
            "Krawiec": "753105",
            "Kucharz": "512001",
            "Kuśnierz": "753106",
            "Lakiernik samochodowy": "713203",
            "Magazynier-logistyk": "432106",
            "Mechanik motocyklowy": "723107",
            "Mechanik pojazdów kolejowych (Opracowano na podstawie dokumentu z dnia 9 kwietnia 2020 r.)": "723318",
            "Mechanik pojazdów samochodowych": "723103",
            "Mechanik precyzyjny": "731103",
            "Mechanik-monter maszyn i urządzeń": "723310",
            "Mechanik-operator maszyn do produkcji drzewnej": "817212",
            "Mechanik-operator pojazdów i maszyn rolniczych": "834103",
            "Modelarz odlewniczy": "721104",
            "Monter budownictwa wodnego": "711701",
            "Monter izolacji budowlanych": "712401",
            "Monter izolacji przemysłowych": "712403",
            "Monter jachtów i łodzi": "711505",
            "Monter kadłubów jednostek pływających": "721406",
            "Monter konstrukcji budowlanych": "711102",
            "Monter nawierzchni kolejowej": "711603",
            "Monter sieci i instalacji sanitarnych": "712618",
            "Monter sieci i urządzeń telekomunikacyjnych": "742202",
            "Monter stolarki budowlanej": "712906",
            "Monter systemów rurociągowych": "712613",
            "Monter zabudowy i robót wykończeniowych w budownictwie": "712905",
            "Murarz-tynkarz": "711204",
            "Obuwnik": "753602",
            "Ogrodnik": "611303",
            "Operator maszyn i urządzeń do przetwórstwa tworzyw sztucznych": "814209",
            "Operator maszyn i urządzeń do robót ziemnych i drogowych": "834209",
            "Operator maszyn i urządzeń odlewniczych": "812107",
            "Operator maszyn i urządzeń przemysłu drzewnego (Wprowadzono na podstawie rozporządzenia z dnia 18 stycznia 2023 r.)": "817213",
            "Operator maszyn i urządzeń przemysłu spożywczego": "816003",
            "Operator maszyn i urządzeń przeróbczych": "811205",
            "Operator maszyn i urządzeń w gospodarce odpadami (Wprowadzono na podstawie rozporządzenia z dnia 18 stycznia 2023 r.)": "313211",
            "Operator maszyn leśnych": "834105",
            "Operator maszyn w przemyśle włókienniczym": "815204",
            "Operator obrabiarek skrawających": "722307",
            "Operator procesów introligatorskich": "732305",
            "Operator urządzeń przemysłu ceramicznego": "818115",
            "Operator urządzeń przemysłu chemicznego": "813134",
            "Operator urządzeń przemysłu szklarskiego": "818116",
            "Opiekun medyczny": "532102",
            "Opiekun osoby starszej": "341202",
            "Opiekun w domu pomocy społecznej": "341203",
            "Opiekunka dziecięca": "325905",
            "Opiekunka środowiskowa": "341204",
            "Optyk-mechanik": "731104",
            "Ortoptystka": "325906",
            "Piekarz": "751204",
            "Podolog (opracowano na podstawie dokumentu z dnia 18 czerwca 2021 r.)": "323014",
            "Pracownik obsługi hotelowej": "962907",
            "Pracownik pomocniczy fryzjera": "932920",
            "Pracownik pomocniczy gastronomii": "941203",
            "Pracownik pomocniczy krawca": "932915",
            "Pracownik pomocniczy mechanika": "932916",
            "Pracownik pomocniczy obsługi hotelowej": "911205",
            "Pracownik pomocniczy stolarza": "932918",
            "Pracownik pomocniczy ślusarza": "932917",
            "Pracownik pomocniczy w gospodarce odpadami (Wprowadzono na podstawie rozporządzenia z dnia 18 stycznia 2023 r.)": "932922",
            "Protetyk słuchu": "321401",
            "Przetwórca mięsa": "751108",
            "Przetwórca ryb": "751103",
            "Pszczelarz": "612302",
            "Rękodzielnik wyrobów włókienniczych": "731808",
            "Rolnik": "613003",
            "Rybak śródlądowy": "622201",
            "Sprzedawca": "522301",
            "Stolarz": "752205",
            "Ślusarz": "722204",
            "Tapicer": "753402",
            "": "N/A"
        }
        # return zawody_dict[specjalnosc]
        return zawody_dict.get(specjalnosc, 'N/A')

    def utworz_wykaz(self):
        szablonWykaz = "Szablony/szablon_wykaz.docx"
        tmp = "Szablony/plik_tymczasowy.docx"
        # wstawianie listy uczniów na końcu dokumentu

        docTempl = docx.Document(szablonWykaz)

        try:
            dfw = pd.read_excel(open(self.plik, "rb"), dtype={'PESEL': str})
            self.wynik.configure(bootstyle="success")

        except Exception as e:

            self.wynik.configure(text=f"Nie wyczytano poprawnyc danych")
            self.wynik.configure(bootstyle="danger")

            return  # Zakończ funkcję, gdy dane nie zostały wczytane poprawnie
        
        if dfw.empty:
            print("Plik nie zawiera danych.")
            # Tutaj możesz dodać obsługę braku danych, np. wyświetlić komunikat użytkownikowi
            return  # Zakończ funkcję, gdy brak danych

        filtered_dfw = dfw[dfw["Dane oddziału"].str.contains(self.var.get(
        ), case=False) & dfw['Specjalność/Zawód'].str.contains(self.combobox.get(), case=False)]

        for linia in range(filtered_dfw.shape[0]):
            rekord = filtered_dfw.iloc[linia].to_dict()
            # new_paragraph = docTempl.add_paragraph(str(linia+1) + ". " + rekord['Imię']+ " " + rekord['Nazwisko'] )
            num_of_paragraphs = len(docTempl.paragraphs)
            # print(num_of_paragraphs)

            npar = docTempl.paragraphs[num_of_paragraphs-3]
            # print(npar)
            npar.add_run(
                (
                    (
                        ((f"{str(linia + 1)}. " + rekord['Imię']) + " ")
                        + rekord['Nazwisko']
                    )
                    + "\n"
                )
            )

        if os.path.exists(tmp):
            os.remove(tmp)
        docTempl.save(tmp)

        # wstawienie danych do szablonu wykazu uczniów

        szablon = DocxTemplate(tmp)
        rekord = filtered_dfw.iloc[linia].to_dict()
        context = {'dataWyst': self.data_wystawienia.entry.get(),
                   'imię': rekord['Imię'],
                   'nazwisko': rekord['Nazwisko'],
                   'dataUrodzenia': rekord['Data urodzenia'],
                   'miejsceUrodzenia': rekord['Miejsce urodzenia'],
                   'PESEL': rekord['PESEL'],
                   'zawod': self.combobox.get(),
                   'kodZawodu': self.symbolZawodu(self.combobox.get()),
                   'dataRozp': self.data_rozpoczecia.entry.get(),
                   'dataZako': self.data_zakonczenia.entry.get(),
                   'godzRozp': self.godzina_rozpoczecia.get()+":"+self.minuty_rozpoczecia.get(),
                   'stopien': self.var.get()
                   }

        # renderowane dokumentu (podstawianie danych ze słownika)
        szablon.render(context)

        if not os.path.exists("../Data"):
            os.mkdir("../Data")

        if not os.path.exists("../Data/Wykazy"):
            os.mkdir("../Data/Wykazy")

        # zapisywanie dokumentu

        szablon.save(
            f"../Data/Wykazy/{rekord['Dane oddziału']}{rekord['Specjalność/Zawód']}.docx")

        # informacja zwrotna

        self.wynik.configure(text=f"utworzono: {str(linia + 1)} pozycji")

        #Ustawienie napisu na przycisku do generowania pdf
        folder_path = os.path.abspath(os.path.join(
            os.path.dirname(__file__), '..', 'Data', 'Wykazy'))
        
        # files_to_convert = os.listdir(folder_path)
        # total_files_wykazy = len(files_to_convert)


        files_to_convert = os.listdir(folder_path)
        total_files_wykazy = sum(1 for file_name in files_to_convert if os.path.isfile(os.path.join(folder_path, file_name)))



        files_processed = 0

        self.btn_utworz_wykaz_pdf.configure(text=f"PDF: {str(total_files_wykazy)} plików")

    def utworz_wykaz_pdf(self):
        folder_path = os.path.abspath(os.path.join(
            os.path.dirname(__file__), '..', 'Data', 'Wykazy'))



        convert(folder_path)

    def utworz_skierowania(self):
        szablon = "Szablony/szablon_skierowanie.docx"
        # tmp = "Szablony/output1.docx"

        # Otwórz plik xlsx


        try:
            df = pd.read_excel(open(self.plik, "rb"), dtype={
                           'PESEL': str, 'Data urodzenia': str})
            self.wynik.configure(bootstyle="success")

        except Exception as e:

            self.wynik.configure(text=f"Nie wyczytano poprawnyc danych")
            self.wynik.configure(bootstyle="danger")

            return  # Zakończ funkcję, gdy dane nie zostały wczytane poprawnie
        
        if df.empty:
            print("Plik nie zawiera danych.")
            # Tutaj możesz dodać obsługę braku danych, np. wyświetlić komunikat użytkownikowi
            return  # Zakończ funkcję, gdy brak danych        
        
        
        
        # df['Data urodzenia'] = pd.to_datetime(
        #     df['Data urodzenia']).dt.strftime('%d.%m.%Y')







        filtered_df = df[df["Dane oddziału"].str.contains(self.var.get(
        ), case=False) & df['Specjalność/Zawód'].str.contains(self.combobox.get(), case=False)]

        # .shape zwraca tupla wiersze, kolumny
        for linia in range(filtered_df.shape[0]):
            rekord = filtered_df.iloc[linia].to_dict()

            doc = DocxTemplate(szablon)

            context = {'dataWyst': self.data_wystawienia.entry.get(),
                       'imię': rekord['Imię'],
                       'nazwisko': rekord['Nazwisko'],
                       'dataUrodzenia': rekord['Data urodzenia'],
                       'miejsceUrodzenia': rekord['Miejsce urodzenia'],
                       'PESEL': rekord['PESEL'],
                       'zawod': self.combobox.get(),
                       'kodZawodu': self.symbolZawodu(self.combobox.get()),
                       'dataRozp': self.data_rozpoczecia.entry.get(),
                       'dataZako': self.data_zakonczenia.entry.get(),
                       'godzRozp': self.godzina_rozpoczecia.get()+":"+self.minuty_rozpoczecia.get(),
                       'stopien': self.var.get()
                       }

            # print(context)

            # renderowane dokumentu (podstawianie danych ze słownika)
            doc.render(context)

            if not os.path.exists("../Data"):
                os.mkdir("../Data")

            if not os.path.exists("../Data/Skierowania"):
                os.mkdir("../Data/Skierowania")

            # zapisywanie dokumentu
            doc.save("../Data/Skierowania/"+rekord['Dane oddziału']+rekord['Specjalność/Zawód'] +
                     rekord['Imię']+rekord['Nazwisko'] + ".docx")

            # informacja zwrotna
            self.wynik.configure(
                text=f"utworzono: {str(linia + 1)} dokumentów")
            
            #Ustawienie napisu na przycisku do generowania pdf
            folder_path_skierowania = os.path.abspath(os.path.join(
                os.path.dirname(__file__), '..', 'Data', 'Skierowania'))
            
            
            # Zliczanie plikó i folderów w folderze Skierowania
            # files_to_convert_skierowania = os.listdir(folder_path_skierowania)
            # total_files_skierowania = len(files_to_convert_skierowania)


            # Zliczanie tylko plików
            files_to_convert_skierowania = os.listdir(folder_path_skierowania)
            total_files_skierowania = sum(1 for file_name in files_to_convert_skierowania if os.path.isfile(os.path.join(folder_path_skierowania, file_name)))





            files_processed = 0

            self.btn_utworz_skierowania_pdf.configure(text=f"PDF: {str(total_files_skierowania)} plików")

    def utworz_skierowania_pdf(self):
        folder_path = os.path.abspath(os.path.join(
            os.path.dirname(__file__), '..', 'Data', 'Skierowania'))
        convert(folder_path)

    def otworz_folder_wykaz(self):
        folder_path = os.path.abspath(os.path.join(
            os.path.dirname(__file__), '..', 'Data', 'Wykazy'))
        subprocess.Popen(f'explorer "{folder_path}"')

    def otworz_folder_skierowania(self):
        folder_path = os.path.abspath(os.path.join(
            os.path.dirname(__file__), '..', 'Data', 'Skierowania'))
        subprocess.Popen(f'explorer "{folder_path}"')

    def credits(self):
        self.pole_tekstowe.delete(1.0, tk.END)
        self.pole_tekstowe.insert(
            tk.END, "Autor: Piotr Dębowski\nZespół Szkół Energetycznych i Usługowych w Łaziskch Górnych\n\nWersja 0.27\n\n")


if __name__ == "__main__":
    app = App()
    app.root.mainloop()
