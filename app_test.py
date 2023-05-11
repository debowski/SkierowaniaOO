import subprocess
import tkinter as tk
import tkinter.ttk as ttk
import ttkbootstrap as ttkb
from tkinter import StringVar
from tkinter import filedialog
import pandas as pd
from docxtpl import DocxTemplate

import docx
import os

class App:

    def __init__(self):

        

        self.root = ttkb.Window(themename="darkly")
        self.root.title("Skierowania 0.22")

        self.root.grid()
        self.root.columnconfigure(0, weight=0, minsize=500)
        self.root.columnconfigure(1, weight=1, minsize=400)
        self.root.rowconfigure(0, weight=1)

        self.dodaj_widzety()

    def dodaj_widzety(self):
        self.frame = ttkb.Frame(self.root)
        self.frame.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)

        self.frame2 = ttkb.Frame(self.root, bootstyle="success")
        self.frame2.grid(row=0, column=1, sticky="nsew", padx=10, pady=10)



        self.frame.columnconfigure(0, weight=1)
        self.frame.columnconfigure(1, weight=1)
        self.frame.rowconfigure(0, weight=0)
        self.frame.rowconfigure(9, weight=1)

        self.frame2.columnconfigure(0, weight=1)
        self.frame2.rowconfigure(0, weight=1)



        self.plik = ""

        self.btn_wyb_plik = ttkb.Button(
            self.frame, text="Wybierz plik", bootstyle="info", command=self.otwarcie_pliku)
        self.btn_wyb_plik.grid(row=0, column=0, sticky="nsew",
                               columnspan=3, padx=5, pady=5)

        # Wybieranie klasy

        self.var = StringVar()

        self.radio1 = ttkb.Radiobutton(
            self.frame, text="Klasa 1", variable=self.var, value="1", bootstyle="success-outline-toolbutton", command=self.set_lista_zawodow_1)
        self.radio2 = ttkb.Radiobutton(
            self.frame, text="Klasa 2", variable=self.var, value="2", bootstyle="warning-outline-toolbutton", command=self.set_lista_zawodow_2)
        self.radio3 = ttkb.Radiobutton(
            self.frame, text="Klasa 3", variable=self.var, value="3", bootstyle="danger-outline-toolbutton", command=self.set_lista_zawodow_3)
        self.radio1.grid(row=1, column=0, sticky="nsew", padx=5, pady=5)
        self.radio2.grid(row=1, column=1, sticky="nsew", padx=5, pady=5)
        self.radio3.grid(row=1, column=2, sticky="nsew", padx=5, pady=5)

        self.radiobuttons = []
        self.lista_zawodow = StringVar()

        self.zawody = ('Sprzedawca',)
        self.zawody1 = ('Wybierz mądrze...',)
        self.zawody2 = ('Wybierz mądrze...',)
        self.zawody3 = ('Wybierz mądrze...',)


        # wstaw combobox z wyborem zawodow

        # Set the default value for the combobox to be the first item in the list of values.
        self.combo_current_var = tk.StringVar()
        self.combobox = ttkb.Combobox(
            self.frame, values=self.zawody, textvariable=self.combo_current_var, bootstyle="success")
        self.combobox.grid(row=2, column=0, sticky="nsew",
                           padx=5, pady=5, columnspan=3)
        self.combobox.configure(state='readonly')
        self.combobox.set(self.zawody[0])
        self.combobox.current(0)

        self.combobox.bind("<<ComboboxSelected>>", self.wypisanie_osob)    

        

        self.lab_data_wystawienia = ttkb.Label(self.frame, text="Data wystawienia")
        self.lab_data_wystawienia.grid(row=3, column=0, sticky="nsew", padx=5, pady=5)


        self.data_wystawienia = ttkb.DateEntry(
            self.frame, firstweekday=0)
        self.data_wystawienia.grid(
            row=3, column=1, sticky="nsew", padx=5, pady=5, columnspan=2)
        
        self.lab_data_rozpoczecia = ttkb.Label(self.frame, text="Data rozpoczęcia")
        self.lab_data_rozpoczecia.grid(row=4, column=0, sticky="nsew", padx=5, pady=5)
        
        
        self.data_rozpoczecia = ttkb.DateEntry(
            self.frame)
        self.data_rozpoczecia.grid(
            row=4, column=1, sticky="nsew", padx=5, pady=5, columnspan=2)
        
        self.lab_data_zakonczenia = ttkb.Label(self.frame, text="Data zakończenia")
        self.lab_data_zakonczenia.grid(row=5, column=0, sticky="nsew", padx=5, pady=5)
        
        
        self.data_zakonczenia = ttkb.DateEntry(
            self.frame)
        self.data_zakonczenia.grid(
            row=5, column=1, sticky="nsew", padx=5, pady=5, columnspan=2)
        



        self.label_godzina_rozpoczecia = ttkb.Label(self.frame, text="Godzina rozpoczęcia")
        self.label_godzina_rozpoczecia.grid(row=6, column=0, sticky="nsew", padx=5, pady=5)

        
        self.godzina_rozpoczecia = ttkb.Spinbox(self.frame, from_=0, to=23, justify="center", format="%02.0f")
        self.godzina_rozpoczecia.grid(row=6, column=1, sticky="nsew", padx=5, pady=5)
        self.godzina_rozpoczecia.insert(0, "08")
        

        self.minuty_rozpoczecia = ttkb.Spinbox(self.frame, from_=0, to=59, justify="center", format="%02.0f")
        self.minuty_rozpoczecia.grid(row=6, column=2, sticky="nsew", padx=5, pady=5)
        self.minuty_rozpoczecia.insert(0, "00")

        # self.btn_test = ttkb.Button(self.frame, text="Test", bootstyle="info", command=self.test)
        # self.btn_test.grid(row=7, column=0, sticky="nsew", padx=5, pady=5, columnspan=3)

        #add 4 buttons
        self.btn_utworz_wykaz = ttkb.Button(self.frame, text="Utwórz wykaz", bootstyle="success", command=self.utworz_wykaz)
        self.btn_utworz_wykaz.grid(row=7, column=0, sticky="nsew", padx=5, pady=5)

        self.btn_otworz_folder_wykaz = ttkb.Button(self.frame, text="Otwórz folder wykaz", command=self.otworz_folder_wykaz)
        self.btn_otworz_folder_wykaz.grid(row=7, column=1, sticky="nsew", padx=5, pady=5, columnspan=2)

        self.btn_utworz_skierowania = ttkb.Button(self.frame, text="Utwórz skierowania", bootstyle="success", command=self.utworz_skierowania)
        self.btn_utworz_skierowania.grid(row=8, column=0, sticky="nsew", padx=5, pady=5)

        self.btn_otworz_folder_skierowania = ttkb.Button(self.frame, text="Otwórz folder skierowania", command=self.otworz_folder_skierowania)
        self.btn_otworz_folder_skierowania.grid(row=8, column=1, sticky="nsew", padx=5, pady=5, columnspan=2)

        
        self.wynik = ttkb.Label(self.frame, text="Wynik", bootstyle="inverse-dark")
        self.wynik.grid(row=9, column=0, sticky="sew", padx=5, pady=5, columnspan=3)
        
        
        # self.separator = ttkb.Separator(self.frame, orient="horizontal", bootstyle="success")
        # self.separator.grid(row=9, column=0, sticky="sew", padx=5, pady=5, columnspan=3)



        # frame2 - przyciski
        self.pole_tekstowe = tk.Text(self.frame2)
        self.pole_tekstowe.grid(row=0, column=0,  padx=1, pady=1, sticky="nsew")
        










    def set_lista_zawodow_1(self) -> None:
        self.combobox['values'] = list(self.zawody1)
        self.wypisanie_osob()

    def set_lista_zawodow_2(self) -> None:
        self.combobox['values'] = list(self.zawody2)
        self.wypisanie_osob()

    def set_lista_zawodow_3(self):
        self.combobox['values'] = list(self.zawody3)
        self.wypisanie_osob()

    def utworz_lz(self):
        df = pd.read_excel(self.plik)
        # print(df.head)

        # Ekstrakcja liczby z kolumny 'Dane oddziału'
        df['Oddział'] = df['Dane oddziału'].str.extract(
            '(\d+)', expand=False).astype(int)

        # Utworzenie trzech zbiorów na podstawie wartości kolumny 'Oddział'
        self.zawody1 = set(df[df['Oddział'] == 1]
                           ['Specjalność/Zawód'].tolist())
        self.zawody2 = set(df[df['Oddział'] == 2]
                           ['Specjalność/Zawód'].tolist())
        self.zawody3 = set(df[df['Oddział'] == 3]
                           ['Specjalność/Zawód'].tolist())

        print(self.zawody1)

    def otwarcie_pliku(self):
        filetypes = (
            ('Arkusze', '*.xlsx'),
            ('All files', '*.*')
        )

        self.plik = filedialog.askopenfilename(title='Wybierz plik',
                                               initialdir='..\\Data',
                                               filetypes=filetypes)

        self.btn_wyb_plik.configure(text=self.plik)
        self.utworz_lz()
        self.radio1.invoke()

    
    def wypisanie_osob(self, event=None):

        if self.plik == "":
            self.brak_pliku()
            return
        else:

            df = pd.read_excel(open(self.plik, "rb"))
            filtered_df = df[df["Dane oddziału"].str.contains(self.var.get(), case=False) & df['Specjalność/Zawód'].str.contains(self.combobox.get(), case=False)]
            tekst = ""
            numer = 1

            for linia in range(filtered_df.shape[0]):
                rekord = filtered_df.iloc[linia].to_dict()

                tekst = tekst + str(numer) + ". " + \
                    rekord['Imię'] + " " + rekord['Nazwisko'] + "\n"

                numer = numer+1

            # wstawianie listy uczniów do ramki prawej
            self.pole_tekstowe.delete(1.0, tk.END)
            self.pole_tekstowe.insert(tk.END, tekst)
        

    def brak_pliku(self):
        self.pole_tekstowe.delete(1.0, tk.END)
        self.pole_tekstowe.insert(tk.END, "Nie wybrano pliku")

    
    def test(self):
        print("działamy")
        self.wypisanie_osob()
        # self.pole_tekstowe.delete(1.0, tk.END)
        # self.pole_tekstowe.insert(tk.END, self.var.get())  

        # self.pole_tekstowe.insert(tk.END, self.data_rozpoczecia.entry.get())
        # self.pole_tekstowe.insert(tk.END, self.combobox.get())


    def symbolZawodu(self, specjalnosc) -> str:

        zawody_dict = {
            'Cukiernik': '751201',
            'Fryzjer': '514101',
            'Sprzedawca': '522301',
            'Mechanik pojazdów samochodowych': '723103',
            'Kucharz': '512001',
            'Blacharz samochodowy': '721306',
            'Piekarz': '751204',
            'Stolarz': '752205',
            'Lakiernik': '713201',
            'Monter sieci, instalacji i urządzeń sanitarnych': '712616',
            'Elektromechanik pojazdów samochodowych': '741203',
            'Elektryk': '741103',
            'Monter zabudowy i robót wykończeniowych w budownictwie': '712905',
            'Mechanik-monter maszyn i urządzeń': '723310',
            'Magazynier-logistyk': '432106',
            '': 'N/A'
        }
        return zawody_dict[specjalnosc]



    def utworz_wykaz(self):
        
        szablon = "Szablony\\szablon.docx"
        szablonWykaz = "Szablony\\szablonWykaz.docx"
        tmp = "Szablony\\output1.docx"
        domyslnyplik = "..\\Data\\WydrukiListXls.xlsx"

        # wstawianie listy uczniów na końcu dokumentu
        
        docTempl = docx.Document(szablonWykaz)
        dfw = pd.read_excel(open(self.plik, "rb"), dtype={'PESEL': str})
        # filtered_dfw = dfw[dfw["Dane oddziału"].str.contains(daneOddzialu(pole_dane_oddzialu
        #                                                                 ), case=False) & dfw['Specjalność/Zawód'].str.contains(lista_specjalnosci.get(), case=False)]

        
        filtered_dfw = dfw[dfw["Dane oddziału"].str.contains(self.var.get(), case=False) & dfw['Specjalność/Zawód'].str.contains(self.combobox.get(), case=False)]

        
        
        
        
        
        for linia in range(filtered_dfw.shape[0]):
            rekord = filtered_dfw.iloc[linia].to_dict()
            # new_paragraph = docTempl.add_paragraph(str(linia+1) + ". " + rekord['Imię']+ " " + rekord['Nazwisko'] )
            num_of_paragraphs = len(docTempl.paragraphs)
            # print(num_of_paragraphs)

            npar = docTempl.paragraphs[num_of_paragraphs-3]
            # print(npar)
            npar.add_run(
                (
                    (
                        ((f"{str(linia + 1)}. " + rekord['Imię']) + " ")
                        + rekord['Nazwisko']
                    )
                    + "\n"
                )
            )

        if os.path.exists(tmp):
            os.remove(tmp)
        docTempl.save(tmp)

        # wstawienie danych do szablonu wykazu uczniów

        szablon = DocxTemplate(tmp)
        rekord = filtered_dfw.iloc[linia].to_dict()
        context = {'dataWyst': self.data_wystawienia.entry.get(),
                'imię': rekord['Imię'],
                'nazwisko': rekord['Nazwisko'],
                'dataUrodzenia': rekord['Data urodzenia'],
                'miejsceUrodzenia': rekord['Miejsce urodzenia'],
                'PESEL': rekord['PESEL'],
                'zawod': self.combobox.get(),
                'kodZawodu': self.symbolZawodu(self.combobox.get()),
                'dataRozp': self.data_rozpoczecia.entry.get(),
                'dataZako': self.data_zakonczenia.entry.get(),
                'godzRozp': self.godzina_rozpoczecia.get()+":"+self.minuty_rozpoczecia.get(),
                'PESEL': rekord['PESEL'],
                'stopien': self.var.get()
                }
        





        # renderowane dokumentu (podstawianie danych ze słownika)
        szablon.render(context)

        if not os.path.exists("..\\Data\\Wykazy"):
            os.mkdir("..\\Data\\Wykazy")

        # zapisywanie dokumentu
        szablon.save("..\\Data\\Wykazy\\"+rekord['Dane oddziału'] +
                    rekord['Specjalność/Zawód'] + ".docx")

        # informacja zwrotna

        self.wynik.configure(text=f"utworzono: {str(linia + 1)} pozycji")























    def utworz_skierowania(self):
        szablon = "Szablony\\szablon.docx"
        szablonWykaz = "Szablony\\szablonWykaz.docx"
        tmp = "Szablony\\output1.docx"
        domyslnyplik = "..\\Data\\WydrukiListXls.xlsx"

        # Otwórz plik xlsx
        df = pd.read_excel(open(self.plik, "rb"), dtype={'PESEL': str})

        filtered_df = df[df["Dane oddziału"].str.contains(self.var.get(), case=False) & df['Specjalność/Zawód'].str.contains(self.combobox.get(), case=False)]





        # .shape zwraca tupla wiersze, kolumny
        for linia in range(filtered_df.shape[0]):
            rekord = filtered_df.iloc[linia].to_dict()

            doc = DocxTemplate(szablon)

            context = {'dataWyst': self.data_wystawienia.entry.get(),
                    'imię': rekord['Imię'],
                    'nazwisko': rekord['Nazwisko'],
                    'dataUrodzenia': rekord['Data urodzenia'],
                    'miejsceUrodzenia': rekord['Miejsce urodzenia'],
                    'PESEL': rekord['PESEL'],
                    'zawod': self.combobox.get(),
                    'kodZawodu': self.symbolZawodu(self.combobox.get()),
                    'dataRozp': self.data_rozpoczecia.entry.get(),
                    'dataZako': self.data_zakonczenia.entry.get(),
                    'godzRozp': self.godzina_rozpoczecia.get()+":"+self.minuty_rozpoczecia.get(),
                    'PESEL': rekord['PESEL'],
                    'stopien': self.var.get()
                    }

            # print(context)

            # renderowane dokumentu (podstawianie danych ze słownika)
            doc.render(context)

            if not os.path.exists("..\\Data\\Skierowania"):
                os.mkdir("..\\Data\\Skierowania")

            # zapisywanie dokumentu
            doc.save("..\\Data\\Skierowania\\"+rekord['Dane oddziału']+rekord['Specjalność/Zawód'] +
                    rekord['Imię']+rekord['Nazwisko'] + ".docx")

            # informacja zwrotna
            self.wynik.configure(text=f"utworzono: {str(linia + 1)} dokumentów")






    def otworz_folder_wykaz(self):
        path = r"..\Data\Wykazy"
        subprocess.Popen(f'explorer "{path}"')

    def otworz_folder_skierowania(self):
        path = r"..\Data\Skierowania"
        subprocess.Popen(f'explorer "{path}"')



if __name__ == "__main__":
    app = App()
    app.root.mainloop()
